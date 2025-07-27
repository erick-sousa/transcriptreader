import streamlit as st
from bs4 import BeautifulSoup
import json
from docx import Document
import requests
import base64
from io import BytesIO

st.set_page_config(page_title="Transcript Cleaner", layout="wide")
st.title("🧼 Transcript Cleaner")
st.caption("Import HTML transcript (e.g. from Descript), clean it up, and export as TXT or DOCX")

# --- Functions ---
def clean_transcript_from_html_json(doc_json):
    try:
        alignment = doc_json["mediaLibrary"]["mediaRefs"][0]["voiceover"]["metadata"]["alignment"]
    except KeyError:
        raise ValueError("Could not find alignment data in the JSON.")

    try:
        voices = doc_json["voices"]
        speaker_names = {v["id"]: v["name"] for v in voices}
    except KeyError:
        speaker_names = {}

    output_lines = []
    current_speaker_id = None
    current_line = []

    for word_obj in alignment:
        word = word_obj.get("word", "")
        speaker_id = word_obj.get("speaker", {}).get("speakerId", "")

        if speaker_id != current_speaker_id:
            if current_line:
                output_lines.append(" ".join(current_line))
                current_line = []
            speaker_name = speaker_names.get(speaker_id, f"Speaker {speaker_id[:6]}")
            output_lines.append("")  # blank line before speaker
            output_lines.append(speaker_name)
            current_speaker_id = speaker_id

        current_line.append(word)

    if current_line:
        output_lines.append(" ".join(current_line))

    return "\n".join(output_lines)

def parse_html(html):
    soup = BeautifulSoup(html, "html.parser")
    doc_json_tag = soup.find("script", {"id": "document", "type": "application/json"})
    if not doc_json_tag:
        raise ValueError("No <script id='document'> JSON block found in HTML.")
    doc_json = json.loads(doc_json_tag.string)
    return clean_transcript_from_html_json(doc_json)

def generate_docx(text):
    buffer = BytesIO()
    doc = Document()
    for line in text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            if line.strip().isupper() or line.strip().istitle():
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
            else:
                doc.add_paragraph(line.strip())
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI Layout ---
html_source = st.radio("Import Method", ["Upload HTML File", "Paste URL"])
html_content = None

if html_source == "Upload HTML File":
    uploaded_file = st.file_uploader("Choose an HTML file", type="html")
    if uploaded_file:
        html_content = uploaded_file.read().decode("utf-8")
else:
    url = st.text_input("Paste public Descript URL")
    if url:
        try:
            response = requests.get(url)
            response.raise_for_status()
            html_content = response.text
        except Exception as e:
            st.error(f"Failed to fetch URL: {e}")

if html_content:
    try:
        transcript = parse_html(html_content)
        st.subheader("📝 Cleaned Transcript Preview")
        st.text_area("Transcript", transcript, height=400)

        col1, col2 = st.columns(2)
        with col1:
            st.download_button("⬇️ Download TXT", transcript, file_name="transcript.txt")
        with col2:
            docx_buf = generate_docx(transcript)
            st.download_button("⬇️ Download DOCX", docx_buf, file_name="transcript.docx")

    except Exception as e:
        st.error(f"Error parsing HTML: {e}")
